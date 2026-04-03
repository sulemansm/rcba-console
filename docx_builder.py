from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io


def build_docx(event, report_text, attendance_details=None, project_details=None):
    """Build a DOCX document with event report, attendance, and project details."""
    
    doc = Document()
    
    # Add logo
    try:
        doc.add_picture("logo.png", width=Inches(1.5))
    except:
        pass  # Logo file may not exist, continue without it
    
    # Add header
    heading = doc.add_heading("Rotaract Club of Bombay Airport", level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Event title
    title_heading = doc.add_heading(event.get("title", "Event Report"), level=1)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Event details section
    doc.add_heading("Event Details", level=2)
    
    if event.get("venue"):
        doc.add_paragraph(f"Venue: {event['venue']}")
    if event.get("start_time"):
        doc.add_paragraph(f"Date: {event['start_time']}")
    if event.get("end_time") and event['end_time'].strip():
        doc.add_paragraph(f"End Time: {event['end_time']}")
    if event.get("chief_guest"):
        doc.add_paragraph(f"Chief Guest: {event['chief_guest']}")
    
    # Add attendance details if provided
    if attendance_details:
        doc.add_heading("Attendance Summary", level=2)
        for key, value in attendance_details.items():
            doc.add_paragraph(f"{key}: {value}")
    
    # Add report content
    doc.add_heading("Event Report", level=2)
    doc.add_paragraph(report_text)
    
    # Add project details if provided
    if project_details:
        doc.add_heading("Project Details", level=2)
        for key, value in project_details.items():
            doc.add_paragraph(f"{key}: {value}")
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer